"""Build the illustrated Chinese thermoelectric empirical-analysis report."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


REPORT_DIR = Path(__file__).resolve().parent
PROJECT_DIR = REPORT_DIR.parent
FIGURE_DIR = PROJECT_DIR / "empirical" / "figures"
OUTPUT_DIR = REPORT_DIR / "output"
OUTPUT_DOCX = OUTPUT_DIR / "zt_thermoelectric_empirical_report.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
NAVY = "203748"
MUTED = "5C6770"
INK = "222222"
LIGHT_FILL = "F2F4F7"
PALE_BLUE = "E8EEF5"
WHITE = "FFFFFF"

WESTERN_FONT = "Calibri"
CJK_FONT = "Microsoft YaHei"


def set_run_font(run, size=None, color=INK, bold=None, italic=None, western=WESTERN_FONT, east_asia=CJK_FONT):
    run.font.name = western
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), western)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), western)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), east_asia)
    run._element.get_or_add_rPr().rFonts.set(qn("w:cs"), western)
    if size is not None:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    lang = run._element.get_or_add_rPr().find(qn("w:lang"))
    if lang is None:
        lang = OxmlElement("w:lang")
        run._element.get_or_add_rPr().append(lang)
    lang.set(qn("w:val"), "zh-CN")
    lang.set(qn("w:eastAsia"), "zh-CN")


def style_paragraph(paragraph, before=0, after=6, line_spacing=1.10, keep_with_next=False):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line_spacing
    fmt.keep_with_next = keep_with_next


def shade_paragraph(paragraph, fill):
    ppr = paragraph._p.get_or_add_pPr()
    shd = ppr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        ppr.append(shd)
    shd.set(qn("w:fill"), fill)


def add_field(paragraph, instruction):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])
    set_run_font(run, size=9, color=MUTED)


def set_cell_shading(cell, fill):
    tcpr = cell._tc.get_or_add_tcPr()
    shd = tcpr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcpr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, bottom=80, start=120, end=120):
    tcpr = cell._tc.get_or_add_tcPr()
    tc_mar = tcpr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tcpr.append(tc_mar)
    for margin, value in (("top", top), ("bottom", bottom), ("start", start), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa):
    tcpr = cell._tc.get_or_add_tcPr()
    tcw = tcpr.find(qn("w:tcW"))
    if tcw is None:
        tcw = OxmlElement("w:tcW")
        tcpr.append(tcw)
    tcw.set(qn("w:w"), str(width_dxa))
    tcw.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tblpr = table._tbl.tblPr
    tblw = tblpr.find(qn("w:tblW"))
    if tblw is None:
        tblw = OxmlElement("w:tblW")
        tblpr.append(tblw)
    tblw.set(qn("w:w"), str(sum(widths)))
    tblw.set(qn("w:type"), "dxa")

    tblind = tblpr.find(qn("w:tblInd"))
    if tblind is None:
        tblind = OxmlElement("w:tblInd")
        tblpr.append(tblind)
    tblind.set(qn("w:w"), "120")
    tblind.set(qn("w:type"), "dxa")

    layout = tblpr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tblpr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        trpr = row._tr.get_or_add_trPr()
        cant_split = OxmlElement("w:cantSplit")
        trpr.append(cant_split)
        for cell, width in zip(row.cells, widths):
            set_cell_width(cell, width)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row):
    trpr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    trpr.append(tbl_header)


def set_cell_text(cell, text, *, bold=False, size=9, color=INK, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    style_paragraph(p, before=0, after=0, line_spacing=1.05)
    run = p.add_run(text)
    set_run_font(run, size=size, color=color, bold=bold)


def add_panel_table(doc, rows):
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    widths = [720, 8640]
    headers = ["子图", "数据与物理解释"]
    for cell, header in zip(table.rows[0].cells, headers):
        set_cell_shading(cell, LIGHT_FILL)
        set_cell_text(cell, header, bold=True, size=9, color=DARK_BLUE, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_repeat_table_header(table.rows[0])
    for panel, explanation in rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], panel, bold=True, size=9, color=BLUE, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(cells[1], explanation, size=8.7)
    set_table_geometry(table, widths)
    return table


def add_caption(doc, text):
    p = doc.add_paragraph(style="Caption")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_paragraph(p, before=3, after=5, line_spacing=1.0, keep_with_next=True)
    run = p.add_run(text)
    set_run_font(run, size=9, color=MUTED, italic=True)
    return p


def add_lead(doc, label, text):
    p = doc.add_paragraph(style="Lead")
    style_paragraph(p, before=0, after=5, line_spacing=1.08, keep_with_next=True)
    label_run = p.add_run(f"{label}  ")
    set_run_font(label_run, size=10.2, color=DARK_BLUE, bold=True)
    text_run = p.add_run(text)
    set_run_font(text_run, size=10.2, color=INK)
    return p


def add_figure_page(doc, number, title, lead, filename, panel_rows, caption):
    doc.add_page_break()
    heading = doc.add_paragraph(style="Heading 1")
    heading.add_run(f"图 {number}　{title}")
    heading.paragraph_format.keep_with_next = True
    add_lead(doc, "本图回答", lead)

    picture_p = doc.add_paragraph()
    picture_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_paragraph(picture_p, before=0, after=0, line_spacing=1.0, keep_with_next=True)
    run = picture_p.add_run()
    inline = run.add_picture(str(FIGURE_DIR / filename), width=Inches(5.95))
    docpr = inline._inline.docPr
    docpr.set("descr", f"图 {number}：{title}")
    add_caption(doc, caption)
    add_panel_table(doc, panel_rows)


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = WESTERN_FONT
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal._element.rPr.rFonts.set(qn("w:ascii"), WESTERN_FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), WESTERN_FONT)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), CJK_FONT)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    heading_specs = {
        "Heading 1": (16, BLUE, 16, 8),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (12, DARK_BLUE, 8, 4),
    }
    for name, (size, color, before, after) in heading_specs.items():
        style = doc.styles[name]
        style.font.name = WESTERN_FONT
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style._element.rPr.rFonts.set(qn("w:ascii"), WESTERN_FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), WESTERN_FONT)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), CJK_FONT)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.10
        style.paragraph_format.keep_with_next = True

    caption = doc.styles["Caption"]
    caption.font.name = WESTERN_FONT
    caption.font.size = Pt(9)
    caption.font.color.rgb = RGBColor.from_string(MUTED)
    caption._element.rPr.rFonts.set(qn("w:eastAsia"), CJK_FONT)

    lead = doc.styles["Intense Quote"] if "Intense Quote" in [s.name for s in doc.styles] else doc.styles["Normal"]
    if lead.name != "Lead":
        try:
            lead = doc.styles.add_style("Lead", 1)
        except ValueError:
            lead = doc.styles["Lead"]
    lead.font.name = WESTERN_FONT
    lead.font.size = Pt(10.2)
    lead.font.color.rgb = RGBColor.from_string(INK)
    lead._element.rPr.rFonts.set(qn("w:eastAsia"), CJK_FONT)
    lead.paragraph_format.space_after = Pt(5)
    lead.paragraph_format.line_spacing = 1.08


def configure_page(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    style_paragraph(hp, before=0, after=0, line_spacing=1.0)
    r = hp.add_run("ZT 总体数据与结构特征")
    set_run_font(r, size=8.5, color=MUTED, bold=True)
    r = hp.add_run("　|　技术报告 · 2026")
    set_run_font(r, size=8.5, color=MUTED)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    style_paragraph(fp, before=0, after=0, line_spacing=1.0)
    r = fp.add_run("第 ")
    set_run_font(r, size=9, color=MUTED)
    add_field(fp, "PAGE")
    r = fp.add_run(" 页")
    set_run_font(r, size=9, color=MUTED)


def add_cover(doc):
    spacer = doc.add_paragraph()
    style_paragraph(spacer, before=0, after=0)
    spacer.paragraph_format.space_before = Pt(92)

    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_paragraph(kicker, before=0, after=16, line_spacing=1.0)
    r = kicker.add_run("热电材料总体数据分析报告")
    set_run_font(r, size=11, color=BLUE, bold=True)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_paragraph(title, before=0, after=10, line_spacing=1.05)
    r = title.add_run("二维热电优值 ZT：\n总体数据分布与结构特征")
    set_run_font(r, size=28, color=NAVY, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_paragraph(subtitle, before=0, after=32, line_spacing=1.15)
    r = subtitle.add_run("真实实验数据、最大成对样本集与结构代理的联合分析")
    set_run_font(r, size=14, color=DARK_BLUE)

    formula = doc.add_paragraph()
    formula.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_paragraph(formula, before=0, after=34, line_spacing=1.0)
    shade_paragraph(formula, PALE_BLUE)
    r = formula.add_run("ZT = S²σT / (κₑ + κL)")
    set_run_font(r, size=17, color=NAVY, bold=True)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_paragraph(meta, before=0, after=6, line_spacing=1.25)
    r = meta.add_run("数据层：Starrydata2 实验曲线 · JARVIS 2D/3D · 项目结构元数据\n")
    set_run_font(r, size=10.5, color=MUTED)
    r = meta.add_run("生成日期：2026 年 8 月 30 日")
    set_run_font(r, size=10.5, color=MUTED)

    summary = doc.add_paragraph()
    summary.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_paragraph(summary, before=28, after=0, line_spacing=1.15)
    r = summary.add_run("20,037 个真实非零 ZT 样品　|　6 张总体图　|　20 个统计面板　|　9 项测试通过")
    set_run_font(r, size=10, color=BLUE, bold=True)


def add_overview(doc):
    doc.add_page_break()
    doc.add_heading("执行摘要与读图方法", level=1)

    p = doc.add_paragraph()
    r = p.add_run("核心结论。")
    set_run_font(r, bold=True, color=DARK_BLUE)
    r = p.add_run(
        " 本报告将原先的单一参数模型图与材料总体统计明确分开。主结果全部来自可配对的实验或固定条件 DFT 数据；"
        "每个关系独立使用其所需字段的最大完整样本集，不要求一个样品同时具备全部物性。"
    )
    set_run_font(r)

    doc.add_heading("数据配对规则", level=2)
    rules = [
        ("真实 ZT", "每个实验 sample_id 只保留一个有效峰值 ZT；其他物性仅从同一样品插值到该峰值温度。"),
        ("电子/热输运", "不要求样品有 ZT，在 300、600、900 K 内对同一样品的两种性质配对。"),
        ("晶格热导温度带", "先对同一样品同温度的重复曲线取中位数，再统计样品分位数，避免曲线采点数造成权重偏差。"),
        ("结构代理", "实验元数据用于相对密度、孔隙率、粒径和形态；JARVIS 用于二维原子几何与弹性声速代理。"),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    for cell, text in zip(table.rows[0].cells, ["分析层", "配对方法"]):
        set_cell_shading(cell, LIGHT_FILL)
        set_cell_text(cell, text, bold=True, size=9.5, color=DARK_BLUE, align=WD_ALIGN_PARAGRAPH.CENTER)
    for label, text in rules:
        cells = table.add_row().cells
        set_cell_text(cells[0], label, bold=True, size=9.3, color=BLUE)
        set_cell_text(cells[1], text, size=9.3)
    set_repeat_table_header(table.rows[0])
    set_table_geometry(table, [1800, 7560])

    doc.add_heading("跨材料候选窗口", level=2)
    add_lead(
        doc,
        "注意",
        "这些范围是高 ZT 样品的总体密集区域，不是对所有化学家族都成立的唯一最优值；最终需要在同一材料体系内联合优化。",
    )
    windows = [
        ("载流子浓度 n", "约 10¹⁹–10²⁰ cm⁻³", "兼顾电导与 Seebeck 的常见高 ZT 区域"),
        ("|S|", "约 150–500 μV/K", "过小则功率因子不足，极大时常伴随低电导"),
        ("电导率 σ", "约 10⁴–10⁶ S/m", "过低限制功率因子，过高会增加 κe"),
        ("晶格热导 κL", "亚 1 至约 1–2 W·m⁻¹·K⁻¹", "低 κL 有利，但必须避免同时破坏电子输运"),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    for cell, text in zip(table.rows[0].cells, ["参数", "候选范围", "物理含义"]):
        set_cell_shading(cell, LIGHT_FILL)
        set_cell_text(cell, text, bold=True, size=9.2, color=DARK_BLUE, align=WD_ALIGN_PARAGRAPH.CENTER)
    for row in windows:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            set_cell_text(cells[i], text, bold=(i == 0), size=9.0, color=BLUE if i == 0 else INK)
    set_repeat_table_header(table.rows[0])
    set_table_geometry(table, [2100, 2600, 4660])

    doc.add_heading("图中统计元素", level=2)
    p = doc.add_paragraph()
    r = p.add_run("六边形颜色")
    set_run_font(r, bold=True, color=DARK_BLUE)
    r = p.add_run("表示区域内样品数的 log10；")
    set_run_font(r)
    r = p.add_run("白色折线")
    set_run_font(r, bold=True, color=DARK_BLUE)
    r = p.add_run("表示横轴分箱后的中位数，浅色带为 25%–75% 分位；")
    set_run_font(r)
    r = p.add_run("Spearman ρ")
    set_run_font(r, bold=True, color=DARK_BLUE)
    r = p.add_run("只描述单调相关，不能直接解释为因果关系。")
    set_run_font(r)


def add_synthesis(doc):
    doc.add_page_break()
    doc.add_heading("结构特征—物理参数映射与设计建议", level=1)
    add_lead(
        doc,
        "总体判断",
        "当前数据库能可靠支持电子输运、热导、相对密度、粒径和原子尺度起伏的总体统计；褶皱几何、孔洞拓扑和模态声子信息仍缺少可直接配对的数据。",
    )

    mapping = [
        ("平面/起伏形状", "面外原子跨度 / √面积", "可表征原子尺度 buckling；不是宏观褶皱幅度"),
        ("骨架/疏松结构", "每原子面内面积", "可作原子网络开放度代理；不等于真实多孔骨架"),
        ("是否多孔", "相对密度推算孔隙率", "可分组描述；原始记录常为阈值或范围"),
        ("组成单元是否很软", "体模量 B 与 √(B/质量密度)", "可作平均声速/软硬代理；跨库有效配对仅 79"),
        ("振动频率与群速度", "当前无直接模态数据", "经验图中不作结论；需 DFPT/Phono3py"),
        ("声学/光学支重叠", "当前无完整分支标签", "只在机制模型层作为待验证假设"),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    for cell, text in zip(table.rows[0].cells, ["结构特征", "当前经验代理", "可解释边界"]):
        set_cell_shading(cell, LIGHT_FILL)
        set_cell_text(cell, text, bold=True, size=9.4, color=DARK_BLUE, align=WD_ALIGN_PARAGRAPH.CENTER)
    for row in mapping:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            set_cell_text(cells[i], text, bold=(i == 0), size=9.1, color=BLUE if i == 0 else INK)
    set_repeat_table_header(table.rows[0])
    set_table_geometry(table, [2300, 2800, 4260])

    doc.add_heading("联合优化而不是单变量极值", level=2)
    paragraphs = [
        (
            "电子侧",
            "先把 n 调入高 ZT 密集区，再通过多谷简并和较大的态密度有效质量维持 |S|，同时避免过大的导电有效质量和强散射损失迁移率。",
        ),
        (
            "晶格侧",
            "目标不是无限降低 κL，而是在降低群速度或寿命的同时，尽量不破坏电导通道。粒径、孔隙和软结构都存在电子—声子双重代价。",
        ),
        (
            "结构侧",
            "原子尺度起伏和开放网络与 σ 的统计关系可作为筛选线索，但必须在同一化学家族、相近温度和相近载流子浓度下复核。",
        ),
    ]
    for label, text in paragraphs:
        p = doc.add_paragraph()
        r = p.add_run(f"{label}：")
        set_run_font(r, bold=True, color=DARK_BLUE)
        r = p.add_run(text)
        set_run_font(r)

    doc.add_heading("报告使用边界", level=2)
    p = doc.add_paragraph()
    r = p.add_run(
        "所有总体关系都混合了不同化学家族、温度、制备状态和文献来源，因此适合发现候选窗口与优先级，"
        "不适合直接给出受控因果系数。模型层的响应面可用于机制推演，但不能代替这里的实验总体分布。"
    )
    set_run_font(r)


def build_report():
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_styles(doc)
    configure_page(doc)
    doc.core_properties.title = "二维热电优值 ZT：总体数据分布与结构特征"
    doc.core_properties.subject = "真实实验数据、最大成对样本集与结构代理的联合分析"
    doc.core_properties.author = "Codex"
    doc.core_properties.keywords = "ZT, thermoelectric, Starrydata2, JARVIS, transport, structure"

    add_cover(doc)
    add_overview(doc)

    figure_specs = [
        (
            1,
            "数据覆盖与最大可用样本集",
            "说明不同物性为什么具有不同样本数，并证明后续每个关系都使用自身最大成对完整集合。",
            "01_data_coverage.png",
            [
                ("a", "直接实验覆盖：Seebeck 33,950、总热导 22,259、真实非零 ZT 20,037；n、μ、κe 数据更稀少。"),
                ("b", "峰值温度配对：ZT–n 440、ZT–S 12,070、ZT–σ 11,738、ZT–κL 3,649；不强求样品具有全部物性。"),
                ("c", "300/600/900 K 热图给出可内插样品数；后续电子图只在同一样品、同一温度内配对。"),
                ("d", "额外层包括相对密度、粒径、弹性–κL 跨库匹配和 JARVIS 2D；褶皱/孔洞拓扑/模态声子标签缺失。"),
            ],
            "图 1　数据覆盖审计。格内数字和横条末端数字是精确样本数。",
        ),
        (
            2,
            "真实实验 ZT 与四个核心物性的总体分布",
            "回答第二张图是否只来自一个材料：不是；每个面板都是全部可配对实验样品，且每个样品只保留峰值 ZT。",
            "02_experimental_ZT_global.png",
            [
                ("a", "ZT–n：N=440，ρ=-0.18。高 ZT 密集于约 10¹⁹–10²⁰ cm⁻³，但跨材料相关很弱，不能理解为 n 越低越好。"),
                ("b", "ZT–|S|：N=12,070，ρ=+0.36。高 ZT 常见于约 150–500 μV/K；特别大的 |S| 往往伴随低 σ。"),
                ("c", "ZT–σ：N=11,738，ρ=+0.28。高 ZT 常见于约 10⁴–10⁶ S/m；更高 σ 同时增加 κe，故不是无限增大。"),
                ("d", "ZT–κL：N=3,649，ρ=-0.44，为四项中最强总体相关。低 κL 有利，但极低 κL 仍不能补偿很差的电子输运。"),
            ],
            "图 2　真实峰值 ZT 总体分布。六边形为样品密度，白线为分箱中位数。",
        ),
        (
            3,
            "全部可用电子输运关系（不要求有 ZT）",
            "利用 300、600、900 K 下同一样品的最大可配对集合，展开 n、S、σ、μ 与功率因子的竞争。",
            "03_electronic_all_available.png",
            [
                ("a", "Pisarenko：952 个温度点/820 个样品。n 增大时 |S| 总体下降；同一 n 的纵向散布反映有效质量、谷简并和能带差异。"),
                ("b", "n–σ：977 点/824 样品。σ 总体随 n 增大，但同一 n 可跨越数个数量级，说明迁移率是决定差异的重要变量。"),
                ("c", "n–μ：831 点/705 样品。高 n 区 μ 总体降低，包含杂质、缺陷和声子散射增强以及材料家族差异。"),
                ("d", "n–PF：912 点/788 样品。低 n 时 σ 不足，高 n 时 |S| 下降，因此形成宽峰/平台而非单调最优。"),
            ],
            "图 3　电子输运总体关系。蓝/橙/粉分别代表 300/600/900 K。",
        ),
        (
            4,
            "全部可用热输运关系与声速代理",
            "用全部实验 κL 曲线、同温度热导配对及弹性跨库匹配，检验晶格与电子热传导的主要控制项。",
            "04_thermal_all_available.png",
            [
                ("a", "κL–T：109,305 个 25 K 网格点/6,317 样品。深/浅带分别为 25–75% 与 10–90%；高温端还受样品覆盖变化影响。"),
                ("b", "κL–κtotal：6,495 点/4,945 样品。虚线为 κL=κtotal；偏离来自 κe、拆分方法与跨曲线插值差异。"),
                ("c", "κe–σ：764 点/582 样品。参考线采用 L=2.0×10⁻⁸ WΩK⁻²，显示高 σ 必然带来更大的电子热漏。"),
                ("d", "√(B/质量密度)–κL：N=79，相关系数≈+0.62。支持更硬、更快晶格通常有更高 κL，但它是平均声速代理，不是模态群速度。"),
            ],
            "图 4　热输运总体关系。面板 d 为按化学式匹配的跨库证据。",
        ),
        (
            5,
            "实验结构元数据与 ZT、κL",
            "把相对密度、孔隙率、晶粒尺寸和报告形态与真实峰值 ZT 或 κL 合并，观察结构统计而非受控因果。",
            "05_structure_metadata_experimental.png",
            [
                ("a", "孔隙率分组：N=1,177。高孔隙组 ZT 较低，说明若孔隙同时严重损失 σ，单纯降低 κL 也不能提高 ZT。"),
                ("b", "粒径分组：N=196。区间间无稳定单调关系；晶界既散射声子也散射载流子，不能据此给出通用最优粒径。"),
                ("c", "样品形态：N=11,288。bulk、晶体、薄膜、粉末等差异还混合了材料家族、测试方向和制备状态。"),
                ("d", "相对密度–κL：N=191，ρ=-0.12，总体很弱。数据集中于 90%–98%，且不少元数据只是阈值/区间。"),
            ],
            "图 5　结构元数据的描述性关联。箱线图红线为中位数，箱体为 25–75%。",
        ),
        (
            6,
            "JARVIS 2D 原子几何与输运（非真实 ZT）",
            "在固定 n=10²⁰ cm⁻³、T=600 K 下，用全部可用二维 DFT 记录检验原子尺度起伏与开放度代理。",
            "06_jarvis2d_shape_transport.png",
            [
                ("a", "面外跨度比–|S|：807 个材料、n/p 两类。相关约 -0.23/-0.14，仅为弱负相关；该量表征 buckling 而非宏观褶皱。"),
                ("b", "面外跨度比–σ：相关约 +0.36/+0.34。提示起伏几何与轨道耦合/色散相关，但不能证明褶皱直接提高电导。"),
                ("c", "每原子面积–|S|：相关约 +0.12/+0.08，关系很弱；面内面积只是原子网络疏密代理。"),
                ("d", "每原子面积–σ：相关约 -0.25/-0.20。较开放网络在固定条件 DFT 数据中电导稍低，但这不等于真实多孔材料。"),
            ],
            "图 6　二维几何代理与电子输运。共 1,609 条 n/p 型记录，不包含实验 ZT。",
        ),
    ]

    for spec in figure_specs:
        add_figure_page(doc, *spec)

    add_synthesis(doc)
    doc.save(OUTPUT_DOCX)
    print(OUTPUT_DOCX)


if __name__ == "__main__":
    build_report()
